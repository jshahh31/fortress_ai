import io
import re
from docx import Document
from docx.shared import Pt
from app.db.store import store
from app.schemas.chat import MessageRole

async def generate_docx(conversation_id: str) -> io.BytesIO | None:
    """
    Fetches the conversation history, finds the last long assistant message (the report),
    and formats it into a DOCX.
    """
    history = await store.get_messages(conversation_id)
    if not history:
        return None

    # Find the last long assistant message (likely the report)
    # Define "long" as > 500 characters for now.
    report_content = None
    for msg in reversed(history):
        if msg.role == MessageRole.ASSISTANT and len(msg.content) > 500:
            report_content = msg.content
            break
    
    if not report_content:
        # Fallback to last assistant message if no long one found
        for msg in reversed(history):
            if msg.role == MessageRole.ASSISTANT:
                report_content = msg.content
                break
                
    if not report_content:
        return None

    doc = Document()
    doc.add_heading('Fortress AI Analysis Report', 0)

    # Basic markdown parsing
    lines = report_content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Headings
        if line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        
        # Lists
        elif line.startswith('* ') or line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\. ', line):
            content = re.sub(r'^\d+\. ', '', line)
            doc.add_paragraph(content, style='List Number')
        
        # Regular paragraph
        else:
            # Handle basic bold/italic inline
            p = doc.add_paragraph()
            _add_formatted_text(p, line)

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def _add_formatted_text(paragraph, text):
    """
    Handles basic bold (**text**) and italic (*text*) markdown.
    """
    # This is a very basic parser and might not handle overlapping or complex markdown
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)
